"""
ai_word_generator.py

Uses GPT to intelligently parse the extracted Markdown into a clean JSON structure,
then renders a professional Word document from that structured data.

Template layout (exact):
  INVESTMENT BUY NOTE   ← centre, bold, underline, Calibri 26pt
  To / From / Date / Subject
  1. Header Information
     Name, Fund Manager, Benchmark, Date, Analyst, NAV (bullets), AUM (bullets)
  2A. Fund Returns   ← returns table
  2B. Detailed Buy Rationale   ← blank dash line
  3. Portfolio & Quantitative Analysis
     Top 5 Stocks table, Top 5 Sectors table, Quantitative Indicators table
  4. Peer Comparison   ← blank dash line
"""

import io
import json
from datetime import datetime
from openai import OpenAI

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Low-level docx helpers ─────────────────────────────────────────────────────

def _cell_shading(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _para(doc, text='', bold=False, size=11, font='Calibri',
          space_before=0, space_after=4, align=None, color=None,
          italic=False, underline=False):
    """Add a simple paragraph and return it."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = underline
        run.font.name = font
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p


def _kv(doc, label: str, value: str, size=11):
    """Bold label followed by plain value on the same line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    lr = p.add_run(f"{label}: ")
    lr.bold = True
    lr.font.name = 'Calibri'
    lr.font.size = Pt(size)
    vr = p.add_run(value or 'Not available')
    vr.font.name = 'Calibri'
    vr.font.size = Pt(size)
    return p


def _bullet(doc, text: str, size=11):
    """Single bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    return p


def _section_heading(doc, text: str, size=13, space_before=14):
    """Bold section heading, same style as numbered headings in the template."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    return p


def _sub_label(doc, text: str, size=11):
    """Bold sub-label like 'Top 5 Stock Holdings:' inside a section."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    return p


def _dash_line(doc, dashes=60):
    """Add a line of dashes as a placeholder for user-filled content."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('-' * dashes)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p


def _table(doc, headers: list, rows: list):
    """Render a data table with a dark navy header row and alternating row shading."""
    if not rows and not headers:
        doc.add_paragraph('Not available').runs[0].font.name = 'Calibri'
        return

    n_cols = max(len(headers), max((len(r) for r in rows), default=1))
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = 'Table Grid'

    # Header
    hdr_cells = tbl.rows[0].cells
    for ci, hdr in enumerate(headers[:n_cols]):
        hdr_cells[ci].text = ''
        p = hdr_cells[ci].paragraphs[0]
        run = p.add_run(str(hdr))
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _cell_shading(hdr_cells[ci], '1E3A6E')

    # Data rows
    for ri, row_data in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        bg = 'EEF2FB' if ri % 2 == 0 else 'FFFFFF'
        for ci in range(n_cols):
            val = str(row_data[ci]) if ci < len(row_data) else ''
            cells[ci].text = ''
            run = cells[ci].paragraphs[0].add_run(val)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            _cell_shading(cells[ci], bg)

    doc.add_paragraph()   # breathing space after table


# ── GPT parsing ────────────────────────────────────────────────────────────────

PARSE_SYSTEM_PROMPT = """\
You are a financial document parser. Given a Markdown mutual fund fact sheet, extract ALL data
into a structured JSON. Follow these rules exactly:
1. Every value must be copied EXACTLY as written — never paraphrase, infer, or omit.
2. If a field is genuinely absent, use null (JSON null, not the string "null").
3. NAV: include the full string as printed, e.g. "₹45.67 (as on April 15, 2026)".
4. AUM: capture EVERY AUM line as a separate item in the array
   (e.g. Month End AUM, Avg. AUM/AAUM). Use the exact label + value from the source.
5. benchmarks: list each benchmark on a separate string, e.g.
   ["Benchmark: Nifty 500 TRI", "Additional Benchmark: Nifty 50 TRI"].
6. exit_load: copy VERBATIM, every word, bracket, date, clause.
7. returns / top5_stocks / top5_sectors / quantitative_indicators:
   extract headers as a list of strings and rows as a list-of-lists of strings.

Return ONLY a raw JSON object (no markdown fences, no commentary) with this exact schema:
{
  "fund_name": "string",
  "fund_manager": "string or null",
  "fund_category": "string or null",
  "nav": ["string", ...],
  "aum": ["string", ...],
  "inception_date": "string or null",
  "benchmarks": ["string", ...],
  "expense_ratio": "string or null",
  "min_investment": "string or null",
  "exit_load": "string or null",
  "lock_in": "string or null",
  "investment_objective": "string or null",
  "quantitative_indicators": {
    "headers": ["string", ...],
    "rows": [["string", ...], ...]
  },
  "returns": {
    "headers": ["string", ...],
    "rows": [["string", ...], ...]
  },
  "top5_stocks": {
    "headers": ["string", ...],
    "rows": [["string", ...], ...]
  },
  "top5_sectors": {
    "headers": ["string", ...],
    "rows": [["string", ...], ...]
  }
}
"""


def _parse_markdown_with_ai(client: OpenAI, markdown_content: str) -> dict:
    """Call GPT-4o to parse the markdown into structured JSON."""
    from modules.ai_extractor import _call_gpt

    user_content = [
        {
            "type": "text",
            "text": (
                "Parse this mutual fund fact sheet markdown into the required JSON structure.\n\n"
                "MARKDOWN:\n---\n"
                f"{markdown_content}\n"
                "---\n\n"
                "Return ONLY a raw JSON object — NO markdown fences, NO explanation."
            )
        }
    ]

    raw = _call_gpt(
        client,
        PARSE_SYSTEM_PROMPT,
        user_content,
        max_tokens=4096,
        model="gpt-4o",
        json_mode=True
    )

    raw = raw.strip()
    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1] if len(parts) > 1 else raw
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.strip()

    return json.loads(raw)


# ── Document builder ────────────────────────────────────────────────────────────

def _build_docx(data: dict) -> io.BytesIO:
    """Build the Word document exactly matching the required template."""
    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.15)
        sec.right_margin  = Inches(1.15)

    # Default body font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    date_str  = datetime.now().strftime("%B %d, %Y")
    fund_name = (data.get("fund_name") or "Fund Name Not Found").strip()

    # ──────────────────────────────────────────────────────────────────────────
    # INVESTMENT BUY NOTE  (centre, bold, underline, Calibri 26pt)
    # ──────────────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after  = Pt(12)
    tr = title_p.add_run("INVESTMENT BUY NOTE")
    tr.bold      = True
    tr.underline = True
    tr.font.name = 'Calibri'
    tr.font.size = Pt(26)

    # ──────────────────────────────────────────────────────────────────────────
    # Memo header block
    # ──────────────────────────────────────────────────────────────────────────
    _kv(doc, "To",      "Investment Committee / Private Banking Clients")
    _kv(doc, "From",    "Senior Investment Analyst")
    _kv(doc, "Date",    date_str)
    _kv(doc, "Subject", f"Investment Recommendation: {fund_name}")

    _para(doc, space_before=6, space_after=6)   # blank spacer

    # ──────────────────────────────────────────────────────────────────────────
    # 1. Header Information
    # ──────────────────────────────────────────────────────────────────────────
    _section_heading(doc, "1. Header Information")

    _kv(doc, "Name of the scheme", fund_name)
    _kv(doc, "Fund Manager",       data.get("fund_manager") or "Not available")

    # Benchmark Index (bullet list)
    _sub_label(doc, "Benchmark Index:")
    benchmarks = data.get("benchmarks") or []
    if benchmarks:
        for bm in benchmarks:
            _bullet(doc, bm)
    else:
        _bullet(doc, "Not available")

    _kv(doc, "Date of Note",   date_str)
    _kv(doc, "Analyst Name",   "AI Investment Analyst")

    # NAV — bullet list (may be multiple plan NAVs)
    _sub_label(doc, "NAV as on Date:")
    nav_list = data.get("nav") or []
    if isinstance(nav_list, str):
        nav_list = [nav_list]
    if nav_list:
        for nav_item in nav_list:
            if nav_item and nav_item.strip():
                _bullet(doc, nav_item.strip())
    else:
        _bullet(doc, "Not available")

    # AUM — bullet list
    _sub_label(doc, "AUM of the scheme:")
    aum_list = data.get("aum") or []
    if aum_list:
        for aum in aum_list:
            if aum and aum.strip():
                _bullet(doc, aum.strip())
    else:
        _bullet(doc, "Not available")

    _para(doc, space_before=6, space_after=6)   # blank spacer

    # ──────────────────────────────────────────────────────────────────────────
    # 2A. Fund Returns
    # ──────────────────────────────────────────────────────────────────────────
    _section_heading(doc, "2A. Fund Returns")

    returns   = data.get("returns") or {}
    r_headers = returns.get("headers") or []
    r_rows    = returns.get("rows")    or []
    _table(doc, r_headers, r_rows)

    _para(doc, space_before=4, space_after=4)

    # ──────────────────────────────────────────────────────────────────────────
    # 2B. Detailed Buy Rationale  (blank — user fills after download)
    # ──────────────────────────────────────────────────────────────────────────
    _section_heading(doc, "2B. Detailed Buy Rationale")
    _dash_line(doc, dashes=10)

    _para(doc, space_before=4, space_after=4)

    # ──────────────────────────────────────────────────────────────────────────
    # 3. Portfolio & Quantitative Analysis
    # ──────────────────────────────────────────────────────────────────────────
    _section_heading(doc, "3. Portfolio & Quantitative Analysis")

    # Top 5 Stock Holdings
    _sub_label(doc, "Top 5 Stock Holdings:")
    stocks   = data.get("top5_stocks") or {}
    s_heads  = stocks.get("headers") or ["Stock Name", "Sector", "% of Portfolio"]
    s_rows   = stocks.get("rows")    or []
    _table(doc, s_heads, s_rows)

    # Top 5 Sector Holdings
    _sub_label(doc, "Top 5 Sector Holdings:")
    sectors  = data.get("top5_sectors") or {}
    sec_hdrs = sectors.get("headers") or ["Sector", "% of Portfolio"]
    sec_rows = sectors.get("rows")    or []
    _table(doc, sec_hdrs, sec_rows)

    # Quantitative Indicators
    _sub_label(doc, "Quantitative Indicators:")
    quant    = data.get("quantitative_indicators") or {}
    q_heads  = quant.get("headers") or ["Metric", "Value"]
    q_rows   = quant.get("rows")    or []
    _table(doc, q_heads, q_rows)

    _para(doc, space_before=4, space_after=4)

    # ──────────────────────────────────────────────────────────────────────────
    # 4. Peer Comparison  (blank — user fills after download)
    # ──────────────────────────────────────────────────────────────────────────
    _section_heading(doc, "4. Peer Comparison")
    _dash_line(doc, dashes=10)

    # ── Stream out ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Public API ─────────────────────────────────────────────────────────────────

def ai_markdown_to_docx(client: OpenAI, markdown_content: str) -> io.BytesIO:
    """
    Entry point called by api.py /download-word.
    1. GPT parses the markdown → structured JSON.
    2. _build_docx renders that JSON → Word document (BytesIO).
    """
    data = _parse_markdown_with_ai(client, markdown_content)
    return _build_docx(data)
