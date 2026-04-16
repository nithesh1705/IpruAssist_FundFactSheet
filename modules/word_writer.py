import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_formatted_text(paragraph, text: str):
    """Parses simple **markdown** and adds it to the paragraph as formatted runs."""
    parts = text.split('**')
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if idx % 2 == 1: # Odd elements are inside ** **
            run.bold = True

def flush_table(doc, table_rows):
    """Creates a Word table from the accumulated rows."""
    if not table_rows:
        return
    num_cols = max((len(r) for r in table_rows), default=0)
    if num_cols > 0:
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        table.style = 'Table Grid'
        for r_idx, row_data in enumerate(table_rows):
            for c_idx, cell_value in enumerate(row_data):
                if c_idx < num_cols:
                    cell = table.cell(r_idx, c_idx)
                    cell.text = "" 
                    p = cell.paragraphs[0]
                    add_formatted_text(p, cell_value)
        doc.add_paragraph()  # add space after table
    table_rows.clear()

def extract_section(markdown_content: str, header: str) -> str:
    """Extracts content under a specific ## header until the next ## or end of string."""
    pattern = rf"##\s+{header}\n(.*?)(?=\n##\s+|\Z)"
    match = re.search(pattern, markdown_content, re.DOTALL | re.IGNORECASE)
    return match.group(1).strip() if match else ""

def get_list_items(text: str) -> list:
    """Extracts all bullet point lines from the text."""
    lines = [line.strip() for line in text.split('\n')]
    results = []
    for line in lines:
        if line.startswith('- ') or line.startswith('* '):
            results.append(line[2:].strip())
    return results

def _render_markdown_block(doc, md_text: str):
    """Renders a block of markdown (like tables or lists) into the doc."""
    in_table = False
    table_rows = []
    
    for line in md_text.split('\n'):
        line = line.strip()
        is_table_line = line.startswith('|') and line.endswith('|')
        
        if (not is_table_line and in_table) or (not line and in_table):
            flush_table(doc, table_rows)
            in_table = False
            
        if not line:
            continue
            
        if is_table_line:
            in_table = True
            row = [cell.strip() for cell in line.split('|')[1:-1]]
            if all(set(c) == {'-'} or set(c) == {':', '-'} or not c for c in row):
                continue
            table_rows.append(row)
            continue
            
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            continue
            
        p = doc.add_paragraph()
        add_formatted_text(p, line)
        
    if in_table and table_rows:
        flush_table(doc, table_rows)

def markdown_to_docx(markdown_content: str, output_path: str = None):
    # Determine the fund name from # Header
    fund_name_match = re.search(r"^#\s+(.*?)$", markdown_content, re.MULTILINE)
    fund_name = fund_name_match.group(1).strip() if fund_name_match else "Fund Name Not Found"
    fund_name = fund_name.replace("**", "")
    
    portfolio_details = extract_section(markdown_content, "Portfolio Details")
    benchmark_idx = extract_section(markdown_content, "Benchmark Index")
    returns_sec = extract_section(markdown_content, "Returns")
    top_5_stock = extract_section(markdown_content, "Top 5 Stock Holdings")
    top_5_sector = extract_section(markdown_content, "Top 5 Sector Holdings")
    quant_ind = extract_section(markdown_content, "Quantitative Indicators")

    doc = Document()
    
    # "INVESTMENT BUY NOTE (this should be center aligned bold, and underligned, calibri font, 26pt)"
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("INVESTMENT BUY NOTE")
    title_run.bold = True
    title_run.underline = True
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(26)
    
    doc.add_paragraph() # spacer
    
    date_str = datetime.now().strftime("%B %d, %Y")
    
    doc.add_paragraph("To: Investment Committee / Private Banking Clients")
    doc.add_paragraph("From: Senior Investment Analyst")
    doc.add_paragraph(f"Date: {date_str}")
    doc.add_paragraph(f"Subject: Investment Recommendation: {fund_name}")
    doc.add_paragraph()
    
    # 1. Header Information
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Header Information")
    h1_run.bold = True
    h1_run.font.size = Pt(14)
    
    fund_manager = "Not available"
    nav = "Not available"
    aum_list = []
    
    for line in portfolio_details.split('\n'):
        l_lower = line.lower()
        if "manager" in l_lower and fund_manager == "Not available":
            # Extract taking care of leading markdown markers like '- ' or '**'
            clean = re.sub(r"^[\-\*\s]+", "", line).replace("**", "")
            fund_manager = re.sub(r"^fund manager\(s\) and managing.since date:?", "", clean, flags=re.IGNORECASE).strip()
            fund_manager = re.sub(r"^fund manager\(s\):?", "", fund_manager, flags=re.IGNORECASE).strip()
            fund_manager = re.sub(r"^fund manager:?", "", fund_manager, flags=re.IGNORECASE).strip()
            if not fund_manager:
                fund_manager = clean
            
        if "nav" in l_lower and nav == "Not available":
            clean = re.sub(r"^[\-\*\s]+", "", line).replace("**", "")
            nav = re.sub(r"^nav(\s?\(.*?\))?:?", "", clean, flags=re.IGNORECASE).strip()
            if not nav:
                nav = clean
                
        if "aum" in l_lower:
            clean = re.sub(r"^[\-\*\s]+", "", line).replace("**", "")
            aum_list.append(clean)
            
    doc.add_paragraph(f"Name of the scheme: {fund_name}")
    doc.add_paragraph(f"Fund Manager: {fund_manager}")
    
    doc.add_paragraph("Benchmark Index:")
    bmarks = get_list_items(benchmark_idx)
    if bmarks:
        for item in bmarks:
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, item)
    else:
        # Sometimes it's just raw text, not bullet points
        lines = [x for x in benchmark_idx.split('\n') if x.strip()]
        for line in lines:
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, line.strip())
            
    doc.add_paragraph(f"Date of Note: {date_str}")
    doc.add_paragraph("Analyst Name: AI Investment Analyst")
    doc.add_paragraph(f"NAV as on Date: {nav}")
    
    doc.add_paragraph("AUM of the scheme:")
    for aum in aum_list:
        p = doc.add_paragraph(style='List Bullet')
        add_formatted_text(p, aum)
        
    doc.add_paragraph()
    
    # 2A. Fund Returns
    h2a = doc.add_paragraph()
    h2a_run = h2a.add_run("2A. Fund Returns")
    h2a_run.bold = True
    h2a_run.font.size = Pt(14)
    if returns_sec:
        _render_markdown_block(doc, returns_sec)
    else:
        doc.add_paragraph("Not available")
    doc.add_paragraph()
    
    # 2B. Detailed Buy Rationale
    h2b = doc.add_paragraph()
    h2b_run = h2b.add_run("2B. Detailed Buy Rationale")
    h2b_run.bold = True
    h2b_run.font.size = Pt(14)
    doc.add_paragraph("----------------------------------------------------------------------")
    doc.add_paragraph()
    
    # 3. Portfolio & Quantitative Analysis
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. Portfolio & Quantitative Analysis")
    h3_run.bold = True
    h3_run.font.size = Pt(14)
    
    p_stocks = doc.add_paragraph()
    p_stocks.add_run("Top 5 Stock Holdings:").bold = True
    if top_5_stock:
        _render_markdown_block(doc, top_5_stock)
    else:
        doc.add_paragraph("Not available")
    # doc.add_paragraph()
        
    p_sectors = doc.add_paragraph()
    p_sectors.add_run("Top 5 Sector Holdings:").bold = True
    if top_5_sector:
        _render_markdown_block(doc, top_5_sector)
    else:
        doc.add_paragraph("Not available")
    # doc.add_paragraph()
        
    p_quant = doc.add_paragraph()
    p_quant.add_run("Quantitative Indicators:").bold = True
    if quant_ind:
        _render_markdown_block(doc, quant_ind)
    else:
        doc.add_paragraph("Not available")
    doc.add_paragraph()
    
    # 4. Peer Comparison
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("4. Peer Comparison")
    h4_run.bold = True
    h4_run.font.size = Pt(14)
    doc.add_paragraph("----------------------------------------------------------------------")

    if output_path:
        doc.save(output_path)
        return output_path
    else:
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
